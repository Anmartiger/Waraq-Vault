"""WaraqVault — end-to-end regression suite (stubbed OCR, throwaway DB).

Covers: format sniffing, batch upload rules, duplicates, job progress and
cancellation, deletion with FTS sync, scoped search, Arabic proclitic recall,
strict short-token matching, language detection and per-format units.

Run:  .venv/Scripts/python.exe tests/test_regression.py   (needs: pip install httpx)
The real waraq.db is never touched."""
import sys, os, io, types, time, tempfile, sqlite3, zipfile
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

PASS, FAIL = 0, []
def check(name, cond, extra=""):
    global PASS
    if cond: PASS += 1; print(f"  OK   {name}")
    else: FAIL.append(name); print(f"  FAIL {name} {extra}")

# ---- stub the OCR engine BEFORE anything imports it (no EasyOCR load) ----
fake = types.ModuleType("engine.ocr_engine")
fake.GPU_AVAILABLE = False
fake.GPU_NAME = None
fake.OCR_DEVICE = "CPU (stub)"
fake.DELAY = 0.0
def _run_ocr(image):
    if fake.DELAY: time.sleep(fake.DELAY)
    return ["سطر عربي من OCR", "stub ocr line"]
fake.run_ocr = _run_ocr
# Geometry-aware entry point. Returning plain strings on purpose: it proves the
# bbox path degrades gracefully to punctuation joining when no coordinates exist.
fake.run_ocr_boxes = _run_ocr
fake.extract_text_from_image = lambda src: " \n ".join(_run_ocr(src))
fake.reader = None
sys.modules["engine.ocr_engine"] = fake

from engine import database
TMPDB = Path(tempfile.gettempdir()) / "waraq_test_problems.db"
if TMPDB.exists(): TMPDB.unlink()
database.DB_PATH = TMPDB

import main
from fastapi.testclient import TestClient
import fitz

PNG = b"\x89PNG\r\n\x1a\n" + b"fakepngdata-A"
PNG2 = b"\x89PNG\r\n\x1a\n" + b"fakepngdata-B"
PNG3 = b"\x89PNG\r\n\x1a\n" + b"fakepngdata-C"

def blank_pdf(pages):
    d = fitz.open()
    for _ in range(pages): d.new_page()
    data = d.tobytes(); d.close(); return data

def wait_job(client, job_id, timeout=30):
    t0 = time.time()
    while time.time() - t0 < timeout:
        j = client.get(f"/jobs/{job_id}").json()
        if j["state"] in ("done", "error", "cancelled"): return j
        time.sleep(0.05)
    raise TimeoutError(job_id)

def upload(client, files, **form):
    payload = [("file", f) for f in files]
    return client.post("/upload", files=payload, data={k: str(v).lower() for k, v in form.items()})

def _zip_nondocx():
    import io, zipfile as zf
    buf = io.BytesIO()
    with zf.ZipFile(buf, "w") as z: z.writestr("data/readme.txt", "hello")
    return buf.getvalue()

client = TestClient(main.app)
with client:
    print("\n=== P4: detect_kind — extension → content-type → magic/content sniffing ===")
    dk = main.detect_kind
    check("pdf by magic, no ext/ctype",  dk("weird", "", b"%PDF-1.7 xx") == "pdf")
    check("png by magic",                dk("noext", "application/octet-stream", PNG) == "image")
    check("jpeg by magic",               dk("noext", "", b"\xff\xd8\xff\xe0junk") == "image")
    check("arabic utf-8 text, no ext",   dk("هندسة_AI", "application/octet-stream", "نص عربي عادي".encode("utf-8")) == "txt")
    check("cp1256 text, no ext",         dk("قديم", "", "ملف عربي قديم".encode("cp1256")) == "txt")
    check("binary w/ NUL rejected",      dk("blob", "application/octet-stream", b"\x01\x02\x00\x04" * 100) is None)
    check("plain zip is NOT docx",       dk("archive", "", __import__("io").BytesIO and _zip_nondocx()) is None)

    print("\n=== upload validation: fail fast before any processing ===")
    r = upload(client, [("empty.png", b"", "image/png")])
    check("0-byte file -> 400", r.status_code == 400, r.text)
    r = upload(client, [(f"i{i}.png", PNG + bytes([i]), "image/png") for i in range(6)])
    check("6 images -> 400 (max 5 images)", r.status_code == 400, r.text)
    r = upload(client, [("mix_img.png", PNG + b"MIX", "image/png"), ("mix_doc.pdf", blank_pdf(1), "application/pdf")])
    check("mixed image+pdf batch now ACCEPTED (202)", r.status_code == 202, r.text)
    if r.status_code == 202:
        j = wait_job(client, r.json()["job_id"])
        check("mixed batch fully indexed", j["state"] == "done" and len(j["result"]["indexed"]) == 2, str(j.get("result")))
    r = upload(client, [("old.doc", b"junk", "application/msword")])
    check(".doc -> 400 with guidance", r.status_code == 400 and "docx" in r.text)

    print("\n=== P4 end-to-end: extensionless Arabic text file becomes searchable ===")
    r = upload(client, [("هندسة_AI", "المرجل يعمل بتقنية الضغط العالي".encode("utf-8"), "application/octet-stream")])
    check("202 + job id", r.status_code == 202 and "job_id" in r.json(), r.text)
    j = wait_job(client, r.json()["job_id"])
    check("job done, file indexed", j["state"] == "done" and j["result"]["indexed"] == ["هندسة_AI"])
    s = client.get("/search", params={"q": "الضغط"}).json()
    check("its content is searchable", s["count"] == 1 and s["results"][0]["filename"] == "هندسة_AI")
    check("stored as txt => unit line", s["results"][0]["unit"] == "line")

    print("\n=== Feature: multi-image batch (3 images) with per-item statuses ===")
    r = upload(client, [("imgA.png", PNG, "image/png"), ("imgB.png", PNG2, "image/png"), ("imgC.png", PNG3, "image/png")])
    check("batch accepted 202", r.status_code == 202, r.text)
    j = wait_job(client, r.json()["job_id"])
    st = [i["status"] for i in j["items"]]
    check("all 3 indexed", j["result"]["indexed"] == ["imgA.png", "imgB.png", "imgC.png"], str(j["result"]))
    check("per-item statuses recorded", st == ["indexed"] * 3, str(st))

    print("\n=== Feature: duplicates — single 409, batch skip, overwrite replaces ===")
    r = upload(client, [("imgA.png", PNG, "image/png")])
    check("single dup -> 409 sync (no OCR wasted)", r.status_code == 409 and r.json()["detail"]["match"] == "content")
    r = upload(client, [("imgA.png", PNG, "image/png"), ("imgD.png", PNG + b"D", "image/png")])
    j = wait_job(client, r.json()["job_id"])
    check("batch dup skipped, new indexed",
          j["result"]["skipped"] == ["imgA.png"] and j["result"]["indexed"] == ["imgD.png"], str(j["result"]))
    check("skipped item carries reason", any(i["status"] == "skipped" and i["detail"] for i in j["items"]))
    r = upload(client, [("imgA.png", PNG, "image/png")], overwrite=True)
    j = wait_job(client, r.json()["job_id"])
    check("overwrite replaces (no duplicate rows)", j["state"] == "done" and j["result"]["replaced"] >= 1, str(j["result"]))
    con = sqlite3.connect(TMPDB)
    n = con.execute("SELECT COUNT(*) FROM documents WHERE filename='imgA.png'").fetchone()[0]
    check("exactly one imgA row remains", n == 1, f"rows={n}")
    con.close()

    print("\n=== P7/P1: image results — OCR runs, unit=block (no fake line numbers) ===")
    s = client.get("/search", params={"q": "stub"}).json()
    img = next(x for x in s["results"] if x["filename"] == "imgA.png")
    check("image unit is block", img["unit"] == "block")
    check("image matches carry no page", all("page" not in m for m in img["matches"]))

    print("\n=== P7: PDF pipeline — real progress events, queue, cancellation ===")
    fake.DELAY = 0.35
    r = upload(client, [("big_scan.pdf", blank_pdf(4), "application/pdf")])
    job_id = r.json()["job_id"]
    saw_pages, saw_progress = False, False
    for _ in range(100):
        j = client.get(f"/jobs/{job_id}").json()
        if j["total_units"] == 4: saw_pages = True
        if j["state"] == "processing" and j["done_units"] >= 1:
            saw_progress = True; break
        time.sleep(0.05)
    check("total = real page count (4)", saw_pages)
    check("done_units advances per completed page", saw_progress)
    rc = client.post(f"/jobs/{job_id}/cancel")
    check("cancel accepted", rc.status_code == 200 and rc.json()["cancelled"] is True)
    j = wait_job(client, job_id)
    check("job ends cancelled", j["state"] == "cancelled")
    docs = client.get("/documents").json()["documents"]
    check("cancelled pdf NOT indexed (no partial rows)", all(d["filename"] != "big_scan.pdf" for d in docs))
    fake.DELAY = 0.0

    print("\n=== P8: force_ocr flag reaches the PDF engine ===")
    calls = {}
    orig = main.process_hybrid_pdf
    def spy(pdf_bytes, ocr_fn, force_ocr=False, pages=None, progress=None, is_cancelled=None):
        calls["force_ocr"] = force_ocr
        return orig(pdf_bytes, ocr_fn, force_ocr=force_ocr, pages=pages,
                    progress=progress, is_cancelled=is_cancelled)
    main.process_hybrid_pdf = spy
    r = upload(client, [("hybrid.pdf", blank_pdf(1), "application/pdf")], force_ocr=True)
    wait_job(client, r.json()["job_id"])
    main.process_hybrid_pdf = orig
    check("force_ocr=True propagated", calls.get("force_ocr") is True)

    print("\n=== Feature: deletion — cascade to FTS, no orphans, 404 on missing ===")
    docs = client.get("/documents").json()["documents"]
    victim = next(d for d in docs if d["filename"] == "imgD.png")
    r = client.delete(f"/documents/{victim['id']}")
    check("delete -> 200", r.status_code == 200 and r.json()["deleted"] == 1)
    check("second delete -> 404", client.delete(f"/documents/{victim['id']}").status_code == 404)
    con = sqlite3.connect(TMPDB)
    nd = con.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    nf = con.execute("SELECT COUNT(*) FROM documents_fts").fetchone()[0]
    con.close()
    check("FTS rows == document rows (no orphans)", nd == nf, f"docs={nd} fts={nf}")
    s = client.get("/search", params={"q": "stub"}).json()
    check("deleted doc gone from search", all(x["filename"] != "imgD.png" for x in s["results"]))

    print("\n=== Feature: scoped search (doc_id filter) ===")
    all_hits = client.get("/search", params={"q": "stub"}).json()
    one = all_hits["results"][0]
    scoped = client.get("/search", params={"q": "stub", "doc_id": one["id"]}).json()
    check("multiple docs unscoped", all_hits["count"] >= 2, str(all_hits["count"]))
    check("scoped -> only the chosen doc", scoped["count"] == 1 and scoped["results"][0]["id"] == one["id"])
    check("clearing scope restores corpus", client.get("/search", params={"q": "stub"}).json()["count"] == all_hits["count"])

    print("\n=== P2: proclitic recall — تخطيطا finds وتخطيطا (no stemmer, no reindex) ===")
    database.insert_document("plan.txt", "text/plain", "كان وتخطيطا للمستقبل واضحا في المشروع", "h-plan")
    res = database.search_documents("تخطيطا")
    hit = next((r_ for r_ in res if r_["filename"] == "plan.txt"), None)
    check("document found via FTS expansion", hit is not None)
    check("وتخطيطا highlighted", hit and "<b>وتخطيطا</b>" in hit["matches"][0]["text"], hit and hit["matches"][0]["text"])
    res = database.search_documents("بالتخطيط")   # user types WITH prefix, doc has bare form?
    # (reverse direction is stemming territory — not claimed; just must not crash)
    check("prefixed query does not crash", isinstance(res, list))

    print("\n=== P3: short tokens are exact — no more false positives/highlights ===")
    database.insert_document("prep1.txt", "text/plain", "ذهبت مع اخي الي السوق", "h-p1")
    database.insert_document("prep2.txt", "text/plain", "هذه المعلومات مفيدة جدا", "h-p2")
    database.insert_document("prep3.txt", "text/plain", "فيلم جميل عن تحفيزه الدائم", "h-p3")
    names = lambda rs: {r_["filename"] for r_ in rs}
    r1 = database.search_documents("مع")
    check("'مع' finds the real preposition", "prep1.txt" in names(r1))
    check("'مع' does NOT drag in المعلومات", "prep2.txt" not in names(r1), str(names(r1)))
    r2 = database.search_documents("في")
    check("'في' does NOT match فيلم/تحفيزه", "prep3.txt" not in names(r2), str(names(r2)))
    hit = next((r_ for r_ in database.search_documents("تخطيطا") if r_["filename"] == "plan.txt"), None)
    check("highlight stays word-bounded", hit and "<b>في</b>" not in hit["matches"][0]["text"])
    r3 = database.search_documents("معلومات")
    check("'معلومات' (≥3) still prefix-matches المعلومات", "prep2.txt" in names(r3))

    print("\n=== P5: real language detection (ar / en / mixed) ===")
    database.insert_document("en.txt", "text/plain", "pure english content about life and money", "h-en")
    database.insert_document("mix.txt", "text/plain", "تقرير المشروع النهائي final project report with english details مكتمل", "h-mx")
    langs = {r_["filename"]: r_["lang"] for r_ in database.search_documents("مشروع") + database.search_documents("english") + database.search_documents("money")}
    check("arabic doc -> ar", langs.get("plan.txt") == "ar", str(langs))
    check("english doc -> en", langs.get("en.txt") == "en", str(langs))
    check("mixed doc -> mixed", langs.get("mix.txt") == "mixed", str(langs))

    print("\n=== P6: DOCX results use paragraph units (¶ in the UI) ===")
    database.insert_document("word.docx", main._DOCX_TYPE, "--- صفحة 1 ---\nفقرة اولى عن الارشيف\nفقرة ثانية", "h-dx")
    hit = next(r_ for r_ in database.search_documents("الارشيف") if r_["filename"] == "word.docx")
    check("docx unit is para", hit["unit"] == "para")
    check("docx keeps its page number", hit["matches"][0].get("page") == 1)

    print("\n=== hardening: quoted/punctuated queries don't break FTS ===")
    check('quoted "مع" behaves like مع', names(database.search_documents('"مع"')) == names(r1))
    check("lone punctuation query -> empty, no crash", database.search_documents('"" !!') == [])

    print("\n=== /status reports the OCR device ===")
    st = client.get("/status").json()
    check("device string present", "CPU (stub)" in st["ingestion_pipeline"], st["ingestion_pipeline"])

    print("\n=== U1: workspaces — grouping, filtering, one-shot group deletion ===")
    r = upload(client, [("w1.txt", "تقرير المشروع الاول".encode("utf-8"), "text/plain")], workspace="مشاريع 2026")
    wait_job(client, r.json()["job_id"])
    r = upload(client, [("w2.txt", "تقرير المشروع الثاني".encode("utf-8"), "text/plain")], workspace="مشاريع 2026")
    wait_job(client, r.json()["job_id"])
    ws = client.get("/workspaces").json()["workspaces"]
    target = next((w for w in ws if w["name"] == "مشاريع 2026"), None)
    check("workspace listed with correct count", target is not None and target["count"] == 2, str(ws))
    docs_ws = client.get("/documents", params={"workspace": "مشاريع 2026"}).json()
    check("/documents?workspace filters", docs_ws["count"] == 2 and
          all(d["workspace"] == "مشاريع 2026" for d in docs_ws["documents"]))
    s = client.get("/search", params={"q": "المشروع", "workspace": "مشاريع 2026"}).json()
    check("search scoped to workspace", s["count"] == 2, str(s["count"]))
    s = client.get("/search", params={"q": "المشروع", "workspace": "لا وجود"}).json()
    check("other workspace -> empty", s["count"] == 0)
    r = client.delete("/workspaces/مشاريع 2026")
    check("workspace deleted in one shot (2 docs)", r.status_code == 200 and r.json()["deleted"] == 2, r.text)
    check("workspace gone from list", all(w["name"] != "مشاريع 2026" for w in client.get("/workspaces").json()["workspaces"]))
    check("missing workspace delete -> 404", client.delete("/workspaces/مشاريع 2026").status_code == 404)

    print("\n=== U1b: bulk deletion of selected files ===")
    ids = []
    for nm in ("b1.txt", "b2.txt", "b3.txt"):
        r = upload(client, [(nm, f"محتوى {nm} للحذف الجماعي".encode("utf-8"), "text/plain")])
        wait_job(client, r.json()["job_id"])
    docs = client.get("/documents").json()["documents"]
    ids = [d["id"] for d in docs if d["filename"] in ("b1.txt", "b2.txt")]
    r = client.post("/documents/delete", json={"ids": ids})
    check("bulk delete -> 2 rows", r.status_code == 200 and r.json()["deleted"] == 2, r.text)
    con = sqlite3.connect(TMPDB)
    nd = con.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    nf = con.execute("SELECT COUNT(*) FROM documents_fts").fetchone()[0]
    con.close()
    check("FTS stays in sync after bulk delete", nd == nf, f"docs={nd} fts={nf}")

    print("\n=== U2: 50-file text batches allowed; Force OCR re-restricts ===")
    many = [(f"t{i}.txt", f"ملف نصي رقم {i} سريع".encode("utf-8"), "text/plain") for i in range(12)]
    r = upload(client, many)
    check("12 text files accepted (no lazy 1-file limit)", r.status_code == 202, r.text)
    j = wait_job(client, r.json()["job_id"])
    check("all 12 indexed", len(j["result"]["indexed"]) == 12, str(len(j["result"]["indexed"])))
    r = upload(client, [(f"t{i}.txt", b"x y z", "text/plain") for i in range(51)])
    check("51 files -> 400", r.status_code == 400)
    r = upload(client, [("f1.txt", b"a b c", "text/plain"), ("f2.txt", b"d e f", "text/plain")], force_ocr=True)
    check("Force OCR + 2 text files -> 400 (strict again)", r.status_code == 400, r.text)
    r = upload(client, [("fo1.png", PNG + b"F1", "image/png"), ("fo2.png", PNG + b"F2", "image/png")], force_ocr=True)
    check("Force OCR + 2 images (<=5) still allowed", r.status_code == 202, r.text)
    wait_job(client, r.json()["job_id"])

    print("\n=== U3/U8: clean index — no page markers in stored text, pages via page_map ===")
    r = upload(client, [("clean3.pdf", blank_pdf(3), "application/pdf")])
    j = wait_job(client, r.json()["job_id"])
    check("3-page scanned pdf indexed", j["state"] == "done", str(j))
    con = sqlite3.connect(TMPDB)
    raw = con.execute("SELECT raw_text, page_map FROM documents WHERE filename='clean3.pdf'").fetchone()
    con.close()
    check("raw_text contains NO page markers", "--- صفحة" not in raw[0], raw[0][:80])
    check("page_map stored instead", raw[1] is not None and "1" in raw[1])
    s = client.get("/search", params={"q": "stub"}).json()
    hit = next(x for x in s["results"] if x["filename"] == "clean3.pdf")
    check("matches carry real pages from page_map",
          [m.get("page") for m in hit["matches"]] == [1, 2, 3], str(hit["matches"]))
    s2 = client.get("/search", params={"q": "صفحة"}).json()
    check("searching the word صفحة finds no injected junk",
          all(x["filename"] != "clean3.pdf" for x in s2.get("results", [])))

    print("\n=== U5: big scanned PDFs need explicit confirmation + page selection ===")
    r = upload(client, [("book12.pdf", blank_pdf(12), "application/pdf")])
    check("12 scanned pages -> 413 confirm_ocr", r.status_code == 413, r.text)
    d = r.json()["detail"]
    check("payload has estimate + device + files", d.get("reason") == "confirm_ocr"
          and d.get("estimate_seconds", 0) > 0 and d.get("files") and d.get("page_selection_allowed") is True, str(d))
    r = upload(client, [("book12.pdf", blank_pdf(12), "application/pdf")], confirmed=True)
    check("confirmed=true proceeds (202)", r.status_code == 202, r.text)
    j = wait_job(client, r.json()["job_id"])
    check("confirmed job completes", j["state"] == "done")
    r = upload(client, [("subset.pdf", blank_pdf(15), "application/pdf")], pages="2-3")
    check("pages=2-3 keeps scanned count under threshold (202)", r.status_code == 202, r.text)
    j = wait_job(client, r.json()["job_id"])
    s = client.get("/search", params={"q": "stub"}).json()
    hit = next(x for x in s["results"] if x["filename"] == "subset.pdf")
    check("only pages 2-3 processed, real page numbers kept",
          hit["match_count"] == 2 and [m.get("page") for m in hit["matches"]] == [2, 3], str(hit["matches"]))
    r = upload(client, [("subset.pdf", blank_pdf(15), "application/pdf")], pages="9-99")
    check("out-of-range pages -> 400", r.status_code == 400, r.text)

    print("\n=== U6: spoofed extensions rejected by content signature ===")
    odt = io.BytesIO()
    with zipfile.ZipFile(odt, "w") as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text")
        z.writestr("content.xml", "<x/>")
    r = upload(client, [("roro.pdf", odt.getvalue(), "application/pdf")])
    check("ODT renamed to .pdf -> 400 naming real type",
          r.status_code == 400 and "ODT" in r.text, r.text[:160])
    con = sqlite3.connect(TMPDB)
    n = con.execute("SELECT COUNT(*) FROM documents WHERE filename='roro.pdf'").fetchone()[0]
    con.close()
    check("no garbage row was inserted", n == 0)
    r = upload(client, [("fake.docx", b"%PDF-1.7 not a docx", _DOCX := "application/vnd.openxmlformats-officedocument.wordprocessingml.document")])
    check("PDF renamed to .docx -> 400", r.status_code == 400)
    r = upload(client, [("bmw.txt", "BMW cars are fast and reliable on highways.".encode("utf-8"), "text/plain")])
    check("text starting with 'BM' NOT misdetected as BMP", r.status_code == 202, r.text)
    wait_job(client, r.json()["job_id"])

    print("\n=== U7: Force OCR reads images embedded in DOCX ===")
    from docx import Document as _Doc
    from PIL import Image as _Img
    png_buf = io.BytesIO()
    _Img.new("RGB", (24, 24), (200, 180, 120)).save(png_buf, format="PNG")
    dx = _Doc()
    dx.add_paragraph("هذا المستند يحتوي صورة مدمجة داخل الملف للتجربة الكاملة")
    dx.add_picture(io.BytesIO(png_buf.getvalue()))
    dxb = io.BytesIO(); dx.save(dxb)
    r = upload(client, [("media.docx", dxb.getvalue(),
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")], force_ocr=True)
    check("docx + force_ocr accepted", r.status_code == 202, r.text)
    j = wait_job(client, r.json()["job_id"])
    con = sqlite3.connect(TMPDB)
    raw = con.execute("SELECT raw_text FROM documents WHERE filename='media.docx'").fetchone()[0]
    con.close()
    check("embedded image text extracted via OCR", "stub ocr line" in raw, raw[:120])
    dx_b = _Doc()
    dx_b.add_paragraph("مستند ثانٍ مختلف المحتوى يحوي صورة مدمجة ايضا للتجربة")
    dx_b.add_picture(io.BytesIO(png_buf.getvalue()))
    dxb2 = io.BytesIO(); dx_b.save(dxb2)
    r = upload(client, [("media2.docx", dxb2.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")])
    j = wait_job(client, r.json()["job_id"])
    con = sqlite3.connect(TMPDB)
    raw2 = con.execute("SELECT raw_text FROM documents WHERE filename='media2.docx'").fetchone()[0]
    con.close()
    check("without force_ocr embedded images are NOT OCRed", "stub ocr line" not in raw2)

    print("\n=== U3b: DOCX headings merge — no jumpy paragraph numbers ===")
    dx2 = _Doc()
    dx2.add_paragraph("تاريخ السيارات")                 # عنوان قصير (كلمتان)
    dx2.add_paragraph("")                               # سطر فارغ
    dx2.add_paragraph("بدأت صناعة السيارات في نهاية القرن التاسع عشر وتطورت بسرعة كبيرة.")
    dx2.add_paragraph("")
    dx2.add_paragraph("محركات حديثة")                   # عنوان قصير آخر
    dx2.add_paragraph("ظهرت المحركات الكهربائية لاحقا وغيرت شكل الصناعة بالكامل.")
    b2 = io.BytesIO(); dx2.save(b2)
    r = upload(client, [("history.docx", b2.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")])
    wait_job(client, r.json()["job_id"])
    con = sqlite3.connect(TMPDB)
    raw3 = con.execute("SELECT raw_text FROM documents WHERE filename='history.docx'").fetchone()[0]
    con.close()
    lines3 = raw3.split("\n")
    check("exactly 2 dense blocks (headings merged, empties dropped)", len(lines3) == 2, str(lines3))
    check("heading merged into its paragraph", lines3[0].startswith("تاريخ السيارات — "), lines3[0][:60])

    print("\n=== U8b: smart merge builds coherent paragraphs from OCR lines ===")
    from engine.textflow import smart_join
    joined = smart_join(["النص الأول يستمر", "حتى نهاية الجملة هنا.", "جملة جديدة تبدأ", "وتنتهي أيضا؟", "بقايا بلا نهاية"])
    check("terminal punctuation closes paragraphs",
          joined == "النص الأول يستمر حتى نهاية الجملة هنا.\nجملة جديدة تبدأ وتنتهي أيضا؟\nبقايا بلا نهاية", joined)

    print("\n=== G1: OCR paragraphs come from box geometry, not punctuation ===")
    from engine.textflow import join_boxes, join_ocr
    def bx(x0, y0, x1, y1, t):
        return ([[x0, y0], [x1, y0], [x1, y1], [x0, y1]], t, 0.9)
    # two boxes on one visual line, then a far-below line = a second paragraph
    boxes = [bx(10, 10, 90, 30, "السطر الأول"), bx(95, 12, 180, 32, "يكمل هنا"),
             bx(10, 40, 120, 60, "سطر ثانٍ في نفس الفقرة"),
             bx(10, 140, 120, 160, "فقرة جديدة بعيدة")]
    out = join_boxes(boxes)
    check("boxes sharing a line are merged", out.splitlines()[0] == "السطر الأول يكمل هنا سطر ثانٍ في نفس الفقرة", out.splitlines()[0])
    check("a large vertical gap starts a new paragraph", len(out.splitlines()) == 2, str(out.splitlines()))
    check("engine order preserved (no re-sorting, RTL safe)",
          out.index("السطر الأول") < out.index("يكمل هنا") < out.index("فقرة جديدة"))
    # no punctuation anywhere: the old heuristic would have produced ONE blob
    check("geometry beats punctuation on unpunctuated text",
          len(out.splitlines()) == 2 and len(smart_join([b[1] for b in boxes]).splitlines()) == 1)
    check("join_ocr still accepts plain strings (fallback intact)",
          join_ocr(["جملة اولى.", "جملة ثانية"]) == "جملة اولى.\nجملة ثانية")
    check("join_ocr handles empty input", join_ocr([]) == "" and join_ocr(None) == "")

    print("\n=== G2: PDF text layer yields real paragraph blocks ===")
    from engine.pdf_engine import page_paragraphs
    d2 = fitz.open(); pg = d2.new_page()
    pg.insert_text((60, 80), "Annual Report", fontsize=20)
    # insert_textbox flows and wraps the text, like a real document paragraph
    pg.insert_textbox(fitz.Rect(60, 110, 330, 230),
                      "The archive contains records of every project we ran during the last "
                      "fiscal year across all departments and regional offices worldwide.", fontsize=11)
    pg.insert_textbox(fitz.Rect(60, 250, 330, 330),
                      "A second, clearly separate paragraph describing the methodology.", fontsize=11)
    buf2 = io.BytesIO(); d2.save(buf2); d2.close()
    doc2 = fitz.open(stream=buf2.getvalue(), filetype="pdf")
    paras = page_paragraphs(doc2[0])
    raw_lines = [l for l in doc2[0].get_text("text").split("\n") if l.strip()]
    doc2.close()
    check("wrapped lines merge into one paragraph",
          any("every project we ran during the last fiscal year" in p for p in paras), str(paras))
    check("fewer blocks than raw newlines", len(paras) < len(raw_lines), f"{len(paras)} vs {len(raw_lines)}")
    check("separate paragraphs stay separate", len(paras) == 3, str(paras))
    check("heading stays its own block", any(p.strip() == "Annual Report" for p in paras), str(paras))

    print("\n=== G3: DOCX headings identified by Word style, not word count ===")
    dx3 = _Doc()
    h = dx3.add_paragraph("A Rather Long Chapter Heading About Cars")   # 7 words
    h.style = dx3.styles["Heading 1"]
    dx3.add_paragraph("بدأت صناعة السيارات في نهاية القرن التاسع عشر وتطورت بسرعة كبيرة جدا.")
    b3 = io.BytesIO(); dx3.save(b3)
    r = upload(client, [("styled.docx", b3.getvalue(), _DOCX)])
    wait_job(client, r.json()["job_id"])
    con = sqlite3.connect(TMPDB)
    raw4 = con.execute("SELECT raw_text FROM documents WHERE filename='styled.docx'").fetchone()[0]
    con.close()
    lines4 = raw4.split("\n")
    check("7-word Heading-styled title still merges", len(lines4) == 1, str(lines4))
    check("merged block keeps both parts",
          lines4[0].startswith("A Rather Long Chapter Heading About Cars — "), lines4[0][:60])
    from engine.docx_engine import _is_heading
    class _FakeStyle:  # a paragraph whose style lookup explodes must not crash
        @property
        def name(self): raise RuntimeError("no style")
    class _FakePara:
        style = _FakeStyle()
    check("style lookup failure falls back to word count",
          _is_heading(_FakePara(), "two words") is True and _is_heading(_FakePara(), "this has four words") is False)

    print("\n=== legacy rows: old marker-based pages still resolve ===")
    database.insert_document("legacy.pdf", "application/pdf",
                             "--- صفحة 7 ---\nسطر قديم يذكر الارشيف هنا", "h-legacy")
    hit = next(r_ for r_ in database.search_documents("الارشيف") if r_["filename"] == "legacy.pdf")
    check("legacy marker still yields page 7", hit["matches"][0].get("page") == 7, str(hit["matches"]))
    check("legacy marker line itself is not a result",
          all("---" not in m["text"] for m in hit["matches"]))

print(f"\n{'='*60}\nPASSED: {PASS}   FAILED: {len(FAIL)}")
for f in FAIL: print("  ✗", f)
if TMPDB.exists(): TMPDB.unlink()
print("temp DB removed; real waraq.db untouched")
sys.exit(1 if FAIL else 0)
